import streamlit as st
import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from io import BytesIO

def half_to_full_width(text):
    trans_table = str.maketrans(
        '0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz.',
        '０１２３４５６７８９ＡＢＣＤＥＦＧＨＩＪＫＬＭＮＯＰＱＲＳＴＵＶＷＸＹＺａｂｃｄｅｆｇｈｉｊｋｌｍｎｏｐｑｒｓｔｕｖｗｘｙｚ．'
    )
    return text.translate(trans_table)

def process_text(content, replacements):
    for pattern, replacement in replacements:
        content = re.sub(pattern, replacement, content, flags=re.MULTILINE)
    return content

def remove_first_duplicate_line(content):
    lines = content.splitlines()
    pattern = r'\b\d{4}\b'
    all_numbers = re.findall(pattern, content)
    count_numbers = {num: all_numbers.count(num) for num in set(all_numbers)}
    first_occurrences = {num: False for num in count_numbers if count_numbers[num] > 1}
    new_lines = []
    for line in lines:
        numbers_in_line = re.findall(pattern, line)
        if any(num in first_occurrences and not first_occurrences[num] for num in numbers_in_line):
            for num in numbers_in_line:
                if num in first_occurrences and not first_occurrences[num]:
                    first_occurrences[num] = True
                    break
            continue
        new_lines.append(line)
    return "\n".join(new_lines)

def normalize_blank_lines(content):
    new_lines = []
    previous_line_was_blank = False
    for line in content.split('\n'):
        if line.strip():
            new_lines.append(line)
            previous_line_was_blank = False
        elif not previous_line_was_blank:
            new_lines.append(line)
            previous_line_was_blank = True
    return "\n".join(new_lines)

def create_document(uploaded_file, output_file):
    content = uploaded_file.getvalue().decode("utf-8")

    # 改行コードを自動変換してファイルを読み取る
    content = content.replace('\r\n', '\n')

    template_file = "temp.docx"

    replacements = [
        #空白の無駄なテキストレイヤー削除
        (r'V\d+, \d+\n{2,}', r''),

        #00;00;47;08 - 00;00;51;03　→　0047 - 0051
        (r'(\d{2});(\d{2});(\d{2});(\d{2})', r'\2\3'),

        #もともとセリフの前にNが入っていた場合、セリフまでの空白含めて削除
        #(r'(^(?:Ｎ|N|N)[\s　]+)(?=.+\n)', r''),

        #2行目以降のセリフの頭に空白があった場合削除
        #(r'(^[\s　]+(?=(?:.+\n)[\s　]*))', r''),

        #例えば1行目と3行目は頭に空白がなく、2行目だけある場合（頭に空白がない行に、空白がある行が挟まれている）など削除
        #(r'^[\s　]+(?=\S+\n+)', r''),

        #V14, 1削除、前半のタイムコードの後ろに　N　セリフ、改行2回して後半のタイムコードの後ろにON
        (r'(\d{4})[\s　]-[\s　](\d{4})\n(V\d{1,2},[\s　]\d)\n((?:.+(?:\n|))*)', r'\1　　N　　\4\n\n\2　　ON\n'),

        #セリフの2行目以降の頭を1行目にそろえる
        (r'(^(?!.*\d{4}(?: |　)*(?:N|ON)(?: |　)*.*).+$)', r'　　　　　　　　　\1'),
        #(r"(^(?!(?:\d{4}　　(?:Ｎ|ＯＮ))).+$)",r"　　　　　　　　　\1")
    ]
    content = process_text(content, replacements)
    content = remove_first_duplicate_line(content)
    content = normalize_blank_lines(content)
    content = half_to_full_width(content)

    doc = Document(template_file)
    doc.add_paragraph(content)

    if not doc.paragraphs[0].text.strip():
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)

    regex = re.compile(r'[０-９]{4}　　ＯＮ')
    for paragraph in doc.paragraphs:
        original_text = paragraph.text
        if regex.search(original_text):
            paragraph.clear()
            last_end = 0
            for match in regex.finditer(original_text):
                paragraph.add_run(original_text[last_end:match.start()])
                highlighted_run = paragraph.add_run(match.group())
                highlighted_run.font.highlight_color = WD_COLOR_INDEX.GRAY_50
                last_end = match.end()
            paragraph.add_run(original_text[last_end:])

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Hiragino Maru Gothic Pro'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Hiragino Maru Gothic Pro')
            run.font.size = Pt(10.5)

    # バイナリ形式でWordファイルを作成
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

st.title("ナレーション原稿作成アプリ")

uploaded_file = st.file_uploader("1. テキストファイルをアップロードしてください", type=["txt"])
if uploaded_file:
    st.write("アップロードが完了しました。")
    output = create_document(uploaded_file, "output.docx")
    st.download_button(
        label="ダウンロード",
        data=output.getvalue(),
        file_name="output.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
